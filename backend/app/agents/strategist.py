"""
agents/strategist.py
Agent C — The Strategist

Always runs autonomously — no HITL gate here.
Drafts GenUI card → auto-publishes to Shopify + Instagram.
Final output appears in:
  - stratpanel (War Room sub-panel)
  - genui panel (Instagram card preview)
"""
from __future__ import annotations

import json

import httpx
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage

from app.core.config import get_settings
from app.core.state_store import push_event, push_sentinel, update_pipeline_status
from app.schemas import (
    AgentRole, GenUICard, GraphUpdatePayload,
    PipelineState, PipelineStatus,
    StreamEvent, StreamEventType,
    StrategistResult,
)

settings = get_settings()

STRATEGIST_SYSTEM_PROMPT = """You are the Strategist Agent in a Market Intelligence War Room.
Draft an Instagram marketing post and a price strategy based on local competitor pricing AND the actual Manufacturer MSRP.

Return ONLY this JSON (no markdown):
{
  "recommendation_type": "instagram_post" | "price_adjustment" | "both",
  "gen_ui_card": {
    "headline": "Short punchy headline ≤10 words",
    "body_copy": "Instagram caption 50-120 words with emojis",
    "cta": "Call-to-action ≤6 words",
    "hashtags": ["#tag1", "#tag2"],
    "suggested_image_prompt": "DALL-E style image prompt",
    "price_adjustment": {
      "product_name": "name of product",
      "old_price": 999.0,
      "new_price": 949.0,
      "reason": "Explain how this undercuts competition while respecting MSRP."
    } or null
  },
  "rationale": "2-3 sentence strategy explanation"
}
Write for a local business in Coimbatore / Tamil Nadu, India.
"""

import os

async def _search_manufacturer_price(mission_id: str, query: str) -> str:
    """Uses Serper API (Google Search) to find manufacturer pricing."""
    serper_api_key = os.environ.get("SERPER_API_KEY")
    if not serper_api_key:
        await _thought(mission_id, "ℹ️ SERPER_API_KEY not found. Simulating manufacturer price search...")
        return "Simulated MSRP: ₹450"
        
    await _thought(mission_id, f"🌐 Searching web for manufacturer price: '{query}'")
    try:
        url = "https://google.serper.dev/search"
        payload = json.dumps({"q": f"{query} manufacturer price MSRP India in INR"})
        headers = {
            'X-API-KEY': serper_api_key,
            'Content-Type': 'application/json'
        }
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.post(url, headers=headers, data=payload)
            results = resp.json()
            snippet = results.get("organic", [{}])[0].get("snippet", "No clear search result.")
            price_guess = [word for word in snippet.split() if '₹' in word or 'Rs' in word]
            price = price_guess[0] if price_guess else "(Price not clearly visible in snippet)"
            
            summary = f"MSRP Search Result: {snippet} | Extracted: {price}"
            await _thought(mission_id, f"💡 Search Engine returned: {summary}")
            return summary
    except Exception as exc:
        await _thought(mission_id, f"⚠️ Search Engine API error: {exc}")
        return "MSRP Search Failed."



async def _thought(mission_id: str, text: str, event: StreamEventType = StreamEventType.THOUGHT) -> None:
    await push_event(StreamEvent(
        event=event, agent=AgentRole.STRATEGIST,
        data=text, mission_id=mission_id,
    ))


async def _graph(mission_id: str, status: str, summary: str = None) -> None:
    payload = GraphUpdatePayload(node_id="strategist", status=status, result_summary=summary)
    await push_event(StreamEvent(
        event=StreamEventType.GRAPH_UPDATE, agent=AgentRole.STRATEGIST,
        data=payload.model_dump_json(), mission_id=mission_id,
    ))


async def _auto_publish(state: PipelineState) -> list[str]:
    """
    Publishing flow — Saves outputs as document files instead of making real API pushes.
    Outputs:
      1. Strategy Docx or MD
      2. Social Posts Text File
    """
    import os
    from datetime import datetime
    
    mission_id = state.mission_id
    result     = state.strategist_result
    log: list[str] = []

    if not result:
        return log

    card = result.gen_ui_card
    tenant_id = getattr(state, "tenant_id", "default_workspace")
    business_name = getattr(state, "_business_name", "Your Business")
    
    output_dir = os.path.join(os.getcwd(), "outputs")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 1. Strategy Doc
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(f'Market Intelligence Strategy: {business_name}', 0)
        doc.add_heading('Recommendation Type', level=1)
        doc.add_paragraph(str(result.recommendation_type))
        doc.add_heading('Rationale', level=1)
        doc.add_paragraph(str(result.rationale))
        
        if card.price_adjustment:
            doc.add_heading('Price Adjustment Required', level=1)
            p = doc.add_paragraph()
            p.add_run(f"Product: ").bold = True
            p.add_run(f"{card.price_adjustment.get('product_name')}\n")
            p.add_run(f"Old Price: ").bold = True
            p.add_run(f"{card.price_adjustment.get('old_price')}\n")
            p.add_run(f"New Price: ").bold = True
            p.add_run(f"{card.price_adjustment.get('new_price')}\n")
            p.add_run(f"Reason: ").bold = True
            p.add_run(f"{card.price_adjustment.get('reason')}\n")
            
        doc_filename = os.path.join(output_dir, f"strategy_{timestamp}.docx")
        doc.save(doc_filename)
        log.append(f"📄 Strategy DOCX generated: outputs/strategy_{timestamp}.docx")
    except Exception as e:
        md_filename = os.path.join(output_dir, f"strategy_{timestamp}.md")
        with open(md_filename, "w", encoding="utf-8") as f:
            f.write(f"# Market Intelligence Strategy: {business_name}\n\n")
            f.write(f"**Recommendation Type:** {result.recommendation_type}\n\n")
            f.write(f"**Rationale:**\n{result.rationale}\n\n")
            if card.price_adjustment:
                f.write("## Price Adjustment\n")
                f.write(f"- Product: {card.price_adjustment.get('product_name')}\n")
                f.write(f"- Old Price: {card.price_adjustment.get('old_price')}\n")
                f.write(f"- New Price: {card.price_adjustment.get('new_price')}\n")
                f.write(f"- Reason: {card.price_adjustment.get('reason')}\n")
        log.append(f"📄 Strategy MD generated: outputs/strategy_{timestamp}.md")

    # 2. Social Media & WhatsApp Posts text file
    social_filename = os.path.join(output_dir, f"social_posts_{timestamp}.txt")
    try:
        caption = (
            f"{card.headline}\n\n{card.body_copy}\n\n{card.cta}\n\n"
            + " ".join(card.hashtags)
        )
        whatsapp_message = f"Market Intelligence War Room 🚀\nMission for {business_name} complete.\nStrategy: {result.recommendation_type}\n"
        if card.price_adjustment:
            whatsapp_message += f"Price updated to ₹{card.price_adjustment.get('new_price')}."
            
        with open(social_filename, "w", encoding="utf-8") as f:
            f.write("===============================\n")
            f.write("      INSTAGRAM POST COPY      \n")
            f.write("===============================\n\n")
            f.write(caption)
            f.write("\n\n")
            f.write("[* Suggested Image Prompt *]\n")
            f.write(card.suggested_image_prompt)
            f.write("\n\n\n")
            f.write("===============================\n")
            f.write("     WHATSAPP NOTIFICATION     \n")
            f.write("===============================\n\n")
            f.write(whatsapp_message)
            f.write("\n")
            
        log.append(f"📱 Social/WhatsApp post copies saved to: outputs/social_posts_{timestamp}.txt")
    except Exception as exc:
        log.append(f"⚠️ Failed to isolate social posts: {exc}")

    return log


async def run_strategist(state: PipelineState) -> PipelineState:
    mission_id = state.mission_id
    analyst    = state.analyst_result

    await _thought(mission_id, "🧠 Strategist online — always autonomous.")
    await _graph(mission_id, "running")

    business_name  = getattr(state, "_business_name", "Your Business")
    niche          = getattr(state, "_niche", "ecommerce")
    city           = getattr(state, "_city", settings.default_city)
    shopify_ids    = getattr(state, "_shopify_product_ids", [])

    # ── Discover Manufacturer Price ───────────────────────────────────────────
    keywords = getattr(state, "_keywords", [])
    msrp_info = "N/A"
    if keywords:
        # Search MSRP for the primary keyword
        msrp_info = await _search_manufacturer_price(mission_id, keywords[0])
    
    context = (
        f"Business: {business_name} | Niche: {niche} | City: {city}\n\n"
        f"Analyst Summary:\n{analyst.summary if analyst else 'N/A'}\n\n"
        + (
            "Key Gaps:\n"
            + "\n".join(f"- [{g.risk_level}] {g.category}: {g.opportunity}"
                        for g in (analyst.gaps if analyst else []))
            or "None identified."
        )
        + f"\n\nManufacturer Web Search Data: {msrp_info}"
        + f"\n\nRecommended price delta: {analyst.recommended_price_delta_pct if analyst else 'N/A'}%"
    )

    llm = ChatGoogleGenerativeAI(
        model=settings.openai_model, temperature=0.7,
        google_api_key=settings.openai_api_key, streaming=False,
    )

    await _thought(mission_id, "✍️  Drafting Instagram copy and pricing strategy...")
    
    try:
        response = await llm.ainvoke([
            SystemMessage(content=STRATEGIST_SYSTEM_PROMPT),
            HumanMessage(content=context),
        ])
        raw: dict = json.loads(response.content)
    except Exception as exc:
        err_msg = str(exc)
        if "API_KEY_INVALID" in err_msg or "INVALID_ARGUMENT" in err_msg or "400" in err_msg:
            await _thought(mission_id, "⚠️ Invalid or missing Gemini API Key. Using mock strategy data for demonstration.", StreamEventType.THOUGHT)
        else:
            await _thought(mission_id, f"⚠️ LLM Strategy Generation Error: {type(exc).__name__} — using fallback.")
            
        raw = {
            "recommendation_type": "instagram_post",
            "gen_ui_card": {
                "headline": "Special Offer Today!",
                "body_copy": "Check out our latest deals — best prices in town. 🎉",
                "cta": "Shop Now",
                "hashtags": [f"#{business_name.replace(' ', '')}"],
                "suggested_image_prompt": "Vibrant local store product banner",
                "price_adjustment": None,
            },
            "rationale": "Fallback strategy due to API validation error.",
        }

    cd = raw["gen_ui_card"]
    gen_ui_card = GenUICard(
        headline              = cd.get("headline", ""),
        body_copy             = cd.get("body_copy", ""),
        cta                   = cd.get("cta", ""),
        hashtags              = cd.get("hashtags", []),
        suggested_image_prompt= cd.get("suggested_image_prompt", ""),
        price_adjustment      = cd.get("price_adjustment"),
    )

    strategist_result = StrategistResult(
        mission_id          = mission_id,
        recommendation_type = raw.get("recommendation_type", "instagram_post"),
        gen_ui_card         = gen_ui_card,
        rationale           = raw.get("rationale", ""),
    )

    await _thought(mission_id, f"📋 Strategy ready: {strategist_result.recommendation_type}")
    await _thought(mission_id, f'💬 Headline: "{gen_ui_card.headline}"')
    await _graph(mission_id, "done", strategist_result.recommendation_type)

    state.strategist_result = strategist_result
    await update_pipeline_status(
        mission_id, PipelineStatus.PUBLISHING, strategist_result=strategist_result
    )

    # ── Auto-publish (no human gate) ──────────────────────────────────────────
    await _thought(mission_id, "🚀 Auto-publishing results...", StreamEventType.PUBLISH_STARTED)
    publish_log = await _auto_publish(state)
    for entry in publish_log:
        await _thought(mission_id, entry)

    await push_event(StreamEvent(
        event=StreamEventType.PUBLISH_COMPLETE,
        agent=AgentRole.STRATEGIST,
        data=json.dumps({"mission_id": mission_id, "log": publish_log}),
        mission_id=mission_id,
    ))

    state.publish_log = publish_log
    state.status = PipelineStatus.COMPLETE
    await update_pipeline_status(
        mission_id, PipelineStatus.COMPLETE, publish_log=publish_log
    )

    await _thought(mission_id, "🏁 Pipeline complete.", StreamEventType.STATUS)
    await push_sentinel(mission_id)
    return state
